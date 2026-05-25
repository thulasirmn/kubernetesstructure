"""
Generate a Word document comparing the existing TRE (App Service) architecture
with the new SRW POC (Kubernetes / AKS) architecture.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, header_color="0F4C81"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Headers
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        set_cell_shading(hdr_cells[i], header_color)

    # Rows
    for r, row_data in enumerate(rows, start=1):
        for c, val in enumerate(row_data):
            cell = table.rows[r].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    return table


# ─────────────────────────────────────────────────────────────────────────────
# Build document
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# Title
title = doc.add_heading(
    'Secure Research Workspace — POC Evaluation',
    level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    'Kubernetes (AKS) vs Azure App Service for Multi-Tenant Research Workspaces')
run.italic = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Meta block
meta_table = doc.add_table(rows=4, cols=2)
meta_table.style = 'Light List'
meta = [
    ('Document', 'POC Architecture Comparison & Migration Plan'),
    ('Author', 'Thulasiraman Subramani — iLink Systems'),
    ('Date', 'May 2026'),
    ('Status', 'Proof of Concept — Validated End-to-End'),
]
for i, (k, v) in enumerate(meta):
    meta_table.rows[i].cells[0].text = k
    meta_table.rows[i].cells[1].text = v
    for run in meta_table.rows[i].cells[0].paragraphs[0].runs:
        run.bold = True

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, '1. Executive Summary', level=1)

add_para(doc,
    'The current Trusted Research Environment (TRE) provisions researcher '
    'workspaces using Azure App Service. While App Service is operationally '
    'simple, it does not scale economically or technically when each workspace '
    'must host multiple containerised analytics applications (Jupyter, RStudio, '
    'custom Docker images) for many concurrent researchers.')

add_para(doc,
    'This Proof of Concept reimplements the workspace provisioning and session '
    'lifecycle on Azure Kubernetes Service (AKS) with an async, Service Bus-driven '
    'control plane. The new architecture delivers fine-grained resource isolation, '
    'true multi-tenancy via Kubernetes namespaces, native container support, and '
    'significantly lower cost at scale.')

add_para(doc,
    'The POC has been validated end-to-end — workspace creation, session launch '
    '(Jupyter and RStudio), Azure Files mount, ingress routing, and asynchronous '
    'background workers are all working against real Azure resources.', italic=True)

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, '2. Problems with the Current App Service-Based TRE Architecture',
            level=1)

problems = [
    ('Container image flexibility',
     'Azure App Service for Containers enforces constraints on image size, startup '
     'duration, and configuration. Hosting arbitrary research images (Jupyter '
     'kernels with custom Python builds, RStudio with R packages, GPU images) is '
     'restrictive. Many community research images need post-startup configuration '
     'that App Service cannot easily accommodate.'),

    ('Cost inefficiency',
     'Each workspace application either consumes its own App Service Plan or shares '
     'one Premium plan with isolation concerns. Plans are billed continuously even '
     'when idle — there is no scale-to-zero on the tiers required for VNet '
     'integration. Cost scales linearly with the number of workspaces, not '
     'with actual usage.'),

    ('Storage mounting limitations',
     'App Service supports a limited Azure Files mount model. Per-user subdirectory '
     'mounts, fine-grained permissions, and dynamic mount targets are not '
     'straightforward. Sharing a single file share across multiple application '
     'instances requires manual configuration.'),

    ('Network isolation at scale',
     'VNet integration is per-plan rather than per-app. Strict workspace isolation '
     'requires either many plans (high cost) or complex NSG / Private Endpoint '
     'topologies. There is no equivalent of a Kubernetes NetworkPolicy to declare '
     'fine-grained inter-app traffic rules.'),

    ('Multi-tenancy at scale',
     'Hard limits on number of apps per plan, no built-in tenant boundary, and '
     'coarse resource quotas. Workspace isolation has to be enforced by '
     'provisioning separate Azure resources for every tenant, multiplying cost '
     'and operational burden.'),

    ('Scaling limitations',
     'Scale-out is per-app, capped by the App Service Plan SKU. Auto-scale rules '
     'are limited compared to Kubernetes Horizontal Pod Autoscaler / Cluster '
     'Autoscaler. There is no ability to oversubscribe compute across many '
     'workspaces while still enforcing per-workspace limits.'),

    ('Resource governance',
     'CPU and memory are sized at the plan SKU level, not per session. Long-running '
     'and burstable workloads cannot be mixed efficiently. GPU workloads are not '
     'supported on standard App Service.'),

    ('Operational complexity',
     'Each app is its own deployment target. Centralised policy enforcement, '
     'rolling updates, and observability require external tooling. There is no '
     'declarative cluster state model akin to Kubernetes manifests.'),
]
for title_text, body in problems:
    p = doc.add_paragraph()
    r = p.add_run(title_text + ' — ')
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run(body)
    r2.font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, '3. POC Architecture — Kubernetes on AKS', level=1)

add_para(doc,
    'The POC is a .NET 8 ASP.NET Core Minimal API following Clean Architecture '
    '(Domain → Core → Infrastructure → API). All long-running Azure / Kubernetes '
    'operations execute asynchronously via Azure Service Bus messages — the HTTP '
    'API returns 202 Accepted immediately and background workers handle the '
    'actual provisioning.')

add_heading(doc, '3.1 Core Components', level=2)
core_components = [
    ('Azure Cosmos DB (NoSQL)',
     'Three containers — workspaces, sessions, secrets. Workspace documents embed '
     'their users and applications. Storage account keys are encrypted at rest '
     'via ASP.NET Core Data Protection.'),
    ('Azure Service Bus (Standard)',
     'Three queues for async provisioning: srw-workspace-provision, '
     'srw-session-stop, srw-workspace-cleanup. Consumers run as BackgroundService '
     'hosts inside the API process.'),
    ('Azure Kubernetes Service (AKS)',
     'One cluster runs both the SRW API and all researcher session pods. Workspace '
     'isolation is enforced by Kubernetes namespaces and a default-deny '
     'NetworkPolicy.'),
    ('Azure Files (per workspace)',
     'One Storage Account + File Share per workspace, mounted into every session '
     'pod via the Azure File CSI driver. All researchers in a workspace share the '
     'same file system.'),
    ('Azure Container Registry',
     'Stores the SRW API image and any custom research images.'),
    ('User-Assigned Managed Identity + Federated Credential',
     'Keyless authentication for the API pod via AKS Workload Identity.'),
    ('ingress-nginx',
     'Path-based routing for session URLs '
     '(http://<host>/s/<sessionSlug>/), WebSocket-enabled for Jupyter/Shiny.'),
]
for name, desc in core_components:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(name + ': ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

add_heading(doc, '3.2 Background Workers (asynchronous control plane)', level=2)
workers = [
    ('WorkspaceProvisioningConsumer', 'Service Bus consumer — ARM-provisions the '
     'Storage Account + File Share, creates the K8s namespace, Secret, and '
     'NetworkPolicy. Max 5 concurrent, 10-min lock renewal.'),
    ('SessionStopConsumer', 'Service Bus consumer — tears down the K8s Deployment, '
     'Service, and Ingress for a session.'),
    ('WorkspaceCleanupConsumer', 'Service Bus consumer — full workspace teardown '
     '(K8s namespace + Azure Storage + Cosmos secret).'),
    ('SessionStatusPoller', 'Periodic worker (15 s) — syncs K8s pod readiness back '
     'into the sessions container in Cosmos.'),
    ('IdleSessionReaper', 'Periodic worker (10 min) — publishes stop messages '
     'for sessions idle longer than the configured threshold (default 8 h).'),
    ('OrphanResourceCleaner', 'Periodic worker (24 h) — reconciles K8s namespaces '
     'against Cosmos workspaces and removes any orphans.'),
]
for name, desc in workers:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(name + ': ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

add_heading(doc, '3.3 Workspace Creation Flow', level=2)
flow_ws = [
    'POST /api/workspaces — API persists the workspace (status=Pending), seeds '
    'two default applications (Jupyter + RStudio), adds the caller as Admin, and '
    'publishes a WorkspaceProvisionMessage to Service Bus.',
    'API returns HTTP 202 Accepted within a few hundred milliseconds.',
    'WorkspaceProvisioningConsumer picks up the message, sets status=Provisioning '
    'in Cosmos.',
    'Azure Storage Account is created via ARM (StorageV2, Standard LRS, '
    'HTTPS-only, TLS 1.2).',
    'Azure File Share is created inside the storage account with the requested '
    'quota in GiB.',
    'The primary storage key is encrypted and written to the Cosmos secrets '
    'container; an azure-storage-creds K8s Secret is created in the workspace '
    'namespace for the CSI driver.',
    'K8s namespace (ws-<name>-<id>) is created with the srw.io/managed-by label '
    'and a default-deny NetworkPolicy.',
    'Cosmos workspace status transitions to Active.',
]
for i, step in enumerate(flow_ws, 1):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(step)
    r.font.size = Pt(11)

add_heading(doc, '3.4 Session Launch Flow', level=2)
flow_sess = [
    'POST /api/workspaces/{id}/sessions — API validates workspace is Active and '
    'the application belongs to the workspace.',
    'Idempotency check — if the user already has a Starting / Running session for '
    'this application, the existing session is returned.',
    'Cosmos session document is created (status=Pending), with deployment, '
    'service, and ingress names derived from a 10-character session slug.',
    'K8s Deployment is created with the application container image, CPU / memory '
    'requests / limits, and the workspace File Share mounted at the application '
    'MountPath via the Azure File CSI driver.',
    'K8s Service (ClusterIP, port 80) is created selecting the session pod.',
    'K8s Ingress rule with the nginx use-regex annotation is created for the '
    'path /s/<slug>/(/|$)(.*).',
    'SessionStatusPoller polls K8s every 15 s; once the deployment has '
    'ReadyReplicas>=1, the session status moves to Running in Cosmos.',
]
for step in flow_sess:
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(step)
    r.font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, '4. Side-by-Side Comparison', level=1)

headers = ['Aspect', 'Existing TRE (App Service)', 'POC (Kubernetes / AKS)']
rows = [
    ['Compute model',
     'One App Service / plan per workspace app',
     'One AKS cluster, one pod per session'],
    ['Container support',
     'Limited (App Service for Containers constraints)',
     'Any Docker image (native K8s)'],
    ['Storage mount',
     'Azure Files via Bring-Your-Own-Storage, limited',
     'Azure Files via CSI driver, full POSIX semantics'],
    ['Isolation boundary',
     'Per App Service Plan + NSG',
     'Per K8s namespace + default-deny NetworkPolicy'],
    ['Multi-tenancy',
     'Per-plan; coarse limits',
     'Per-namespace; native RBAC and ResourceQuota'],
    ['Resource control',
     'Plan SKU (coarse)',
     'Per-pod requests and limits (fine)'],
    ['Scaling',
     'Per-app, capped by plan SKU',
     'HPA / VPA / Cluster Autoscaler'],
    ['Idle cost',
     'Always-on plans (no scale-to-zero on Premium V3)',
     'Cluster Autoscaler removes idle nodes'],
    ['Async provisioning',
     'Custom orchestration per app',
     'Service Bus + Background workers'],
    ['Networking',
     'VNet integration per plan',
     'Native CNI + NetworkPolicy + ingress controller'],
    ['Custom domains / TLS',
     'Per-app cert config',
     'ingress-nginx + cert-manager (one-time setup)'],
    ['Observability',
     'Per-app App Insights',
     'Cluster-wide via Container Insights + Prometheus'],
    ['Operational model',
     'Manage many apps individually',
     'Declarative manifests, kubectl, GitOps-friendly'],
    ['GPU workloads',
     'Not supported on App Service',
     'AKS GPU node pools'],
    ['Cost at 50 workspaces',
     'Linear: ~50 plans / many Premium tiers',
     'Sub-linear: shared cluster, autoscaling'],
]
add_table(doc, headers, rows)

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, '5. Resource Footprint at Scale (50 Workspaces × 10 Users)',
            level=1)
add_para(doc,
    'The POC architecture has been sized for the target scale of 50 workspaces, '
    'each with up to 10 researchers running both Jupyter and RStudio '
    'concurrently. Despite supporting 1,000 concurrent session pods, the '
    'control-plane footprint stays small.')

footprint_rows = [
    ['Azure Storage Accounts', '50 (one per workspace)'],
    ['Azure File Shares', '50 (one per workspace)'],
    ['Cosmos DB Account / Database', '1 / 1 (three containers)'],
    ['Service Bus Namespace / Queues', '1 / 3'],
    ['AKS Clusters', '1 (shared platform)'],
    ['K8s Namespaces', '52 (50 workspaces + platform + ingress-nginx)'],
    ['K8s Deployments (at full load)', '1,000'],
    ['K8s Services / Ingress rules (at full load)', '1,000 each'],
    ['Storage Account quota (default 250 per subscription)', 'Within limits'],
]
add_table(doc, ['Resource', 'Count / Notes'], footprint_rows)

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, '6. Validated End-to-End', level=1)
validated = [
    'Workspace provisioning saga completes from HTTP request to '
    'status=Active in Cosmos (Service Bus → ARM → AKS).',
    'Storage Account, File Share, K8s Namespace, Secret, NetworkPolicy all '
    'created in the correct subscription and resource group.',
    'Session launch creates Deployment + Service + Ingress; pod mounts the '
    'workspace File Share via Azure File CSI driver.',
    'Bidirectional sync verified: a file created in the pod appears in the '
    'Azure File Share within seconds, and vice-versa.',
    'ingress-nginx path-based routing works for both Jupyter and RStudio '
    '(use-regex annotation required for the /s/<slug>/(/|$)(.*) pattern).',
    'SessionStatusPoller correctly updates session status from Starting → '
    'Running once ReadyReplicas≥1.',
    'Async DELETE flow: workspace and session deletes return 202 immediately, '
    'background consumers tear down all dependent resources.',
]
for v in validated:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(v)
    r.font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, '7. Migration Plan', level=1)

phases = [
    ('Phase 1 — POC validation (Complete)',
     ['Implement async, Service Bus-driven control plane in .NET 8.',
      'Stand up AKS, Cosmos DB, Service Bus, ACR, Managed Identity, RBAC.',
      'Validate Jupyter + RStudio session launch with shared Azure Files mount.',
      'Single-script provisioning (setup-azure.sh / .ps1) for the entire footprint.']),

    ('Phase 2 — Production hardening',
     ['Replace Service Bus connection string with Workload Identity '
      '(Service Bus Data Owner RBAC for the API pod identity).',
      'Replace Cosmos DB account key with Cosmos DB Built-in Data Contributor '
      'role on the pod identity.',
      'Restore PublicNetworkAccess=Disabled on workspace storage accounts and '
      'attach the AKS subnet to the storage IP rules / use Private Endpoints.',
      'Custom domain + cert-manager + Let\'s Encrypt for HTTPS on session URLs.',
      'Replace X-User-Id header with Keycloak OIDC bearer-token authentication.',
      'Apply per-workspace ResourceQuota and LimitRange in K8s.']),

    ('Phase 3 — Observability and operations',
     ['Container Insights enabled on AKS.',
      'Application Insights instrumentation for the SRW API.',
      'Service Bus dead-letter queue alerts.',
      'Cosmos DB throughput / RU monitoring.',
      'Grafana dashboards for cluster, ingress, and per-workspace usage.']),

    ('Phase 4 — Migration of existing workspaces',
     ['Catalog existing App Service-backed workspaces and their data.',
      'Mirror Azure Files content to the new per-workspace storage accounts.',
      'Re-provision workspaces in the new platform with users and applications '
      'pre-seeded from the existing catalog.',
      'Run both platforms side-by-side; cut over per workspace once researchers '
      'are migrated.',
      'Decommission old App Service Plans after stabilisation period.']),

    ('Phase 5 — Optional extensions',
     ['Custom application catalog (researchers can register their own Docker '
      'images per workspace).',
      'GPU node pool for ML workloads.',
      'Per-user data isolation (re-introducing subPath if compliance demands).',
      'Cross-region replication for disaster recovery.',
      'Cost reporting per workspace via Cosmos session history.']),
]

for phase_name, items in phases:
    p = doc.add_paragraph()
    r = p.add_run(phase_name)
    r.bold = True
    r.font.size = Pt(12)
    for it in items:
        bp = doc.add_paragraph(style='List Bullet')
        br = bp.add_run(it)
        br.font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, '8. Risks and Mitigations', level=1)

risk_rows = [
    ['Kubernetes operational learning curve',
     'Provide runbooks; use managed AKS with auto-upgrade; centralise '
     'orchestration in the SRW API so most operations are abstracted from '
     'researchers.'],
    ['Storage account quota (250 per subscription per region)',
     'Sufficient headroom for 50 workspaces; request quota increase if '
     'planning beyond ~200 workspaces or co-locate workspaces in a single '
     'storage account.'],
    ['Service Bus message redelivery / poison messages',
     'MaxDeliveryCount=10 with dead-letter queue; alerting on dead-letter '
     'depth; consumers are idempotent and skip already-Active workspaces.'],
    ['Public IP exposure for ingress',
     'Use internal load balancer or Application Gateway with WAF; private '
     'cluster option available for high-compliance deployments.'],
    ['Concurrent provisioning errors',
     'ARM long-running operations isolated to background workers '
     '(MaxConcurrentCalls configurable); 10-minute lock renewal covers '
     'slow ARM calls; saga retries on transient failure.'],
    ['Multi-subscription / multi-tenant identity',
     'Azure:SubscriptionId is explicit in configuration to avoid the '
     '"default subscription" pitfall observed in this POC.'],
]
add_table(doc, ['Risk', 'Mitigation'], risk_rows)

# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, '9. Recommendation', level=1)
add_para(doc,
    'The Kubernetes-based POC architecture comprehensively addresses every '
    'limitation identified in the existing App Service-based TRE. It offers '
    'lower per-workspace cost at scale, native multi-tenancy, true container '
    'flexibility, and a clean asynchronous control plane that decouples HTTP '
    'requests from long-running Azure provisioning.')

add_para(doc,
    'It is recommended to proceed with the production hardening phases outlined '
    'above and plan a phased migration of existing workspaces. The POC code, '
    'manifests, and provisioning scripts are immediately reusable.', bold=True)

doc.add_paragraph()
add_heading(doc, 'Appendix — Repository Layout', level=1)
appendix = [
    'src/SRW.Domain — Entities only (Workspace, UserSession, '
    'WorkspaceApplication, WorkspaceUser).',
    'src/SRW.Core — Application services, abstractions, '
    'WorkspaceProvisioningService, SessionLauncher.',
    'src/SRW.Infrastructure — Azure ARM, Kubernetes client, Cosmos '
    'repositories, Service Bus messaging, Background workers.',
    'src/SRW.Api — ASP.NET Core Minimal API endpoints, DI composition, '
    'health endpoint.',
    'k8s/manifests — 00-cluster-setup.yaml (RBAC), 10-api-deployment.yaml.',
    'scripts/setup-azure.sh and setup-azure.ps1 — idempotent Azure '
    'provisioning scripts.',
    'AZURE_SETUP.md — step-by-step CLI runbook.',
    'RESOURCE_TOPOLOGY.md — full Azure + K8s resource breakdown.',
]
for it in appendix:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(it)
    r.font.size = Pt(11)

# Save
output = 'SRW-POC-Architecture-Comparison.docx'
doc.save(output)
print(f'OK — wrote {output}')
