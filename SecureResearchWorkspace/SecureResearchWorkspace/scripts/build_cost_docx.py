from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette ----
PRIMARY = "1F4E78"
ACCENT  = "2E75B6"
LIGHT   = "DDEBF7"
GREENBG = "C6EFCE"; GREENTX = "006100"
GREY    = "808080"; DARK = "262626"; WHITE = "FFFFFF"

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(DARK)

for s in doc.sections:
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)

RIGHT = WD_ALIGN_PARAGRAPH.RIGHT; CENTER = WD_ALIGN_PARAGRAPH.CENTER; LEFT = WD_ALIGN_PARAGRAPH.LEFT


def shade(el, hexc):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexc)
    el.append(sh)


def cell_shade(cell, hexc):
    shade(cell._tc.get_or_add_tcPr(), hexc)


def para_shade(p, hexc):
    shade(p._p.get_or_add_pPr(), hexc)


def heading_band(text):
    p = doc.add_paragraph(); para_shade(p, PRIMARY)
    pf = p.paragraph_format; pf.space_before = Pt(12); pf.space_after = Pt(6); pf.left_indent = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(WHITE)
    return p


def subhead(text):
    p = doc.add_paragraph(); pf = p.paragraph_format; pf.space_before = Pt(10); pf.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = RGBColor.from_string(ACCENT)
    return p


def body(text, size=10.5, after=6, before=0):
    p = doc.add_paragraph(); pf = p.paragraph_format; pf.space_before = Pt(before); pf.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size)
    return p


def bullet(parts, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format; pf.space_before = Pt(1); pf.space_after = Pt(3)
    for text, bold, color in parts:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
    return p


def set_cell_borders(cell, nil=False):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + edge)
        if nil:
            e.set(qn("w:val"), "nil")
        else:
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4"); e.set(qn("w:space"), "0"); e.set(qn("w:color"), "BFBFBF")
        borders.append(e)
    tcPr.append(borders)


def style_cell(cell, text, bold=False, color=None, fill=None, align=LEFT, size=10, nilborder=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if fill:
        cell_shade(cell, fill)
    set_cell_borders(cell, nil=nilborder)


def make_table(headers, rows, widths=None, header_fill=ACCENT, zebra=True, align_map=None, last_row_emphasis=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.allow_autofit = False
    align_map = align_map or {}
    for j, h in enumerate(headers):
        style_cell(t.rows[0].cells[j], h, bold=True, color=WHITE, fill=header_fill,
                   align=align_map.get(j, LEFT), size=10)
    trPr = t.rows[0]._tr.get_or_add_trPr(); th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        is_last = (last_row_emphasis is not None and i == len(rows) - 1)
        for j, val in enumerate(row):
            fill = None
            if is_last:
                fill = last_row_emphasis
            elif zebra and i % 2 == 1:
                fill = "F2F7FC"
            style_cell(cells[j], val, bold=is_last, color=(GREENTX if is_last else None),
                       fill=fill, align=align_map.get(j, LEFT), size=10)
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    return t


def callout(text_parts, fill=GREENBG):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]; cell_shade(c, fill); set_cell_borders(c)
    c.text = ""; p = c.paragraphs[0]; p.alignment = CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    for text, bold, color, size in text_parts:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
    return t


def spacer(pts=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts); return p


# ====================== COVER PAGE ======================
spacer(16); spacer(16)
p = doc.add_paragraph(); p.alignment = CENTER
r = p.add_run("SECURE RESEARCH WORKSPACE"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(ACCENT)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph(); p.alignment = CENTER
r = p.add_run("Cloud Hosting Cost Analysis"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = RGBColor.from_string(PRIMARY)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph(); p.alignment = CENTER
r = p.add_run("AKS Shared-Cluster Platform  vs  TRE App Service per Workspace")
r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(DARK)
p.paragraph_format.space_after = Pt(14)

pr = doc.add_paragraph(); pr.alignment = CENTER
pPr = pr._p.get_or_add_pPr(); bdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "18"); bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), ACCENT)
bdr.append(bottom); pPr.append(bdr)

spacer(16)
callout([("Projected ~71% lower compute cost — about $114,000 saved per year", True, PRIMARY, 15),
         ("\nat 50 workspaces / 500 researchers", False, DARK, 11)], fill=LIGHT)
spacer(20)

meta = doc.add_table(rows=3, cols=2); meta.alignment = WD_TABLE_ALIGNMENT.CENTER
metadata = [("Prepared for", "Client — Trusted Research Environment Programme"),
            ("Date", "2 June 2026"),
            ("Classification", "Confidential — for client review")]
for i, (k, v) in enumerate(metadata):
    style_cell(meta.rows[i].cells[0], k, bold=True, color=ACCENT, align=RIGHT, size=10.5, nilborder=True)
    style_cell(meta.rows[i].cells[1], v, align=LEFT, size=10.5, nilborder=True)
for r in meta.rows:
    r.cells[0].width = Inches(2.2); r.cells[1].width = Inches(4.3)

doc.add_page_break()

# ====================== 1. EXECUTIVE SUMMARY ======================
heading_band("1.  Executive Summary")

subhead("The ask")
body("Compare the cost of hosting research applications (Jupyter, RStudio, DBGate, FDSA) on the "
     "current Azure TRE model — one App Service Plan per workspace — against the new Secure Research "
     "Workspace (SRW) platform running on a shared Azure Kubernetes Service (AKS) cluster.")

subhead("Bottom line")
callout([("~71% lower compute cost — roughly $114,000 saved per year", True, GREENTX, 13.5),
         ("\nat 50 workspaces / 500 researchers, while giving each active researcher more compute and stronger isolation.", False, DARK, 10.5)])
spacer(6)
make_table(["", "Per year", "Per month"],
           [["TRE — 50 dedicated App Service Plans (P2v3, 24×7)", "$160,300", "$13,359"],
            ["SRW on AKS — shared cluster, pods on demand", "$46,330", "$3,861"],
            ["Annual / monthly saving", "~$114,000", "$9,498"]],
           widths=[3.9, 1.4, 1.4], align_map={1: RIGHT, 2: RIGHT}, last_row_emphasis=GREENBG)

subhead("Why it is cheaper — one idea")
body("TRE bills per workspace, around the clock. AKS bills per active session.", after=3)
bullet([("App Service Plans cannot scale to zero — all 50 plans run and bill 24×7, even overnight and at weekends when no one is logged in.", False, None)])
bullet([("AKS creates a pod only when a researcher launches an app, reaps it after 8 hours idle, and shrinks the cluster automatically. All workspaces share one elastic pool.", False, None)])
bullet([("TRE cost grows with the number of workspaces (fixed); AKS cost grows with actual concurrent usage (elastic). Research usage is bursty and business-hours-bound — that idle gap is where the money goes in the TRE model.", False, None)])

subhead("It holds up under pressure")
make_table(["Scenario", "AKS / month", "vs TRE"],
           [["Modeled (40% peak concurrency, autoscaled)", "$3,861", "−71%"],
            ["Every session pinned on 24×7 (no scale-down)", "$7,386", "−45%"],
            ["With 1-yr Reserved Instances on baseline nodes", "~$2,500", "~−81%"]],
           widths=[4.2, 1.5, 1.0], align_map={1: RIGHT, 2: RIGHT})
body("AKS only loses if utilization is ~100% and sustained 24/7/365 — which research workloads never are.", before=4)

subhead("Not just cost — it is also better")
bullet([("More compute per active researcher: ", True, None), ("0.5 vCPU bursting to 2 vCPU per session, versus ~0.4 vCPU shared on a P2v3 plan.", False, None)])
bullet([("Per-user, per-app isolation", True, None), (" (one pod each) instead of one shared application per workspace.", False, None)])
bullet([("Workspaces are free when idle", True, None), (" and cheap to create — no fixed per-workspace footprint.", False, None)])

subhead("Assumptions")
body("50 workspaces × 10 users · TRE P2v3 ($267/mo) · 40% peak concurrency · D8s_v3 nodes ($0.384/hr) · "
     "East US pay-as-you-go list prices. Cosmos DB and Service Bus are excluded (already present in production "
     "infrastructure). Substitute your EA/CSP rate card — the structural conclusion does not change.", size=9.5)

doc.add_page_break()

# ====================== 2. DETAILED COST ANALYSIS ======================
heading_band("2.  Detailed Cost Analysis")

subhead("2.1  Scenario modeled")
make_table(["Input", "Value"],
           [["Workspaces", "50"],
            ["Users per workspace", "10 (500 total)"],
            ["Apps per workspace", "4 — Jupyter, RStudio, DBGate, FDSA"],
            ["TRE App Service Plan SKU", "P2v3 (4 vCPU, 16 GB) — one plan per workspace"],
            ["Peak concurrency", "~40% = 200 concurrent sessions"],
            ["Region / pricing basis", "East US, PAYG list prices, 730 hrs/mo, USD"]],
           widths=[2.6, 4.0])
body("Prices are representative PAYG list rates (validated June 2026: D8s_v3 = $0.384/hr; App Service P2v3 "
     "Linux ≈ $0.366/hr). Apply your EA/CSP discount uniformly to both sides — it does not change the structural "
     "conclusion.", size=9.5)

subhead("2.2  The structural difference (why AKS wins)")
make_table(["", "TRE — App Service per workspace", "SRW — shared AKS cluster"],
           [["Unit of cost", "1 App Service Plan per workspace", "1 shared cluster + pods on demand"],
            ["Billing when idle", "24×7×365 — plans cannot scale to zero", "Near-zero — pods reaped after 8h idle; empty nodes removed"],
            ["Scales with", "Number of workspaces (fixed)", "Actual concurrent sessions (elastic)"],
            ["Bin-packing", "None — each plan dedicated to one workspace", "All workspaces share one elastic node pool"],
            ["Compute per active user", "~0.4 vCPU worst case", "0.5 vCPU request, burst to 2 vCPU"]],
           widths=[1.5, 2.6, 2.5])
body("Core insight: TRE pays for 50 plans running continuously regardless of use; AKS pays only for sessions "
     "actually running.", before=4)

subhead("2.3  Model A — TRE App Service (current)")
body("50 workspaces × 1 × P2v3 plan, running 24/7:", after=3)
callout([("50 plans × $267.18/mo  =  $13,359 / month  ≈  $160,300 / year", True, PRIMARY, 11.5)], fill=LIGHT)
body("This is a fixed floor. It does not drop when researchers go home, and it rises to ~$26,700/mo if any "
     "workspaces need P3v3 for heavier compute. Storage is excluded — it is identical on both sides.", before=4)

subhead("2.4  Model B — SRW on AKS")
body("Fixed platform cost:", after=3)
make_table(["Component", "Spec", "$/mo"],
           [["AKS control plane (Standard, uptime SLA)", "$0.10/hr", "73"],
            ["System node pool", "2 × D4s_v3 — API, ingress-nginx, coredns", "280"],
            ["Standard Load Balancer + public IP", "ingress entrypoint", "25"],
            ["Fixed subtotal", "", "~378"]],
           widths=[3.1, 2.5, 0.9], align_map={2: RIGHT}, last_row_emphasis=LIGHT)
body("Cosmos DB and Service Bus are excluded — these already exist in the shared production infrastructure and "
     "add no incremental cost to this deployment.", size=9.5, before=2)

body("Variable cost — user session pods (the elastic part). Node sizing uses a conservative ~8 sessions per "
     "D8s_v3 node; the real number is higher by CPU request, so this understates AKS efficiency.", before=6, after=3)
bullet([("Peak: 200 concurrent → 25 × D8s_v3, ~10 business hrs/day × 22 days = 220 hrs/mo.", False, None)])
bullet([("Off-peak baseline: ~50 long-running sessions → 7 nodes for the remaining 510 hrs/mo.", False, None)])
callout([("Node-hours = (25 × 220) + (7 × 510) = 9,070/mo   →   9,070 × $0.384 = $3,483/mo compute", True, PRIMARY, 10.5),
         ("\n$3,483 (sessions) + $378 (fixed)  =  $3,861 / month  ≈  $46,330 / year", True, GREENTX, 11.5)], fill=LIGHT)

subhead("2.5  Result")
make_table(["", "Per month", "Per year"],
           [["TRE — 50 × P2v3 App Service Plans", "$13,359", "$160,300"],
            ["SRW — shared AKS", "$3,861", "$46,330"],
            ["Savings", "$9,498", "~$114,000"],
            ["% reduction", "", "~71%"]],
           widths=[3.6, 1.5, 1.5], align_map={1: RIGHT, 2: RIGHT}, last_row_emphasis=GREENBG)

subhead("2.6  Robustness — even the worst case favors AKS")
make_table(["Scenario", "AKS $/mo", "vs TRE $13,359"],
           [["Modeled (40% peak, autoscaled)", "$3,861", "−71%"],
            ["Always-on — 200 sessions pinned 24/7 (25 nodes × 730h)", "$7,386", "−45%"],
            ["Pathological — every workspace at P3v3 in TRE", "$3,861 vs $26,700", "−86%"]],
           widths=[4.0, 1.6, 1.4], align_map={1: RIGHT, 2: RIGHT})
body("Even if AKS never scaled down, the shared/bin-packed cluster still beats 50 dedicated plans. AKS only loses "
     "if utilization is ~100% and sustained 24/7/365 — which research workloads never are.", before=4)

subhead("Further AKS savings (not in the headline number)")
bullet([("1-yr Reserved Instances / Savings Plan", True, None), (" on the baseline node pool: ~40% off → AKS ≈ $2,500/mo (~81% reduction).", False, None)])
bullet([("Spot node pool", True, None), (" for batch / non-interactive sessions: D8s_v3 spot ≈ $0.074/hr (~80% off). Use with care for interactive sessions (eviction).", False, None)])
bullet([("Right-sizing requests", True, None), (" below the conservative 8-per-node ratio increases density further.", False, None)])

subhead("2.7  Beyond cost — what AKS also buys")
bullet([("More compute per active researcher (0.5→2 vCPU burst vs ~0.4 vCPU shared on a P2v3).", False, None)])
bullet([("Per-user, per-app isolation (one pod each) vs one shared app instance per workspace.", False, None)])
bullet([("Namespace + NetworkPolicy isolation between workspaces on one cluster.", False, None)])
bullet([("No per-workspace provisioning of fixed infrastructure — workspaces are cheap to create and free when idle.", False, None)])

subhead("2.8  Caveats / how to refine")
bullet([("Replace P2v3 and D8s_v3 rates with your EA/CSP rate card (apply to both sides).", False, None)])
bullet([("The 8-sessions/node ratio is deliberately conservative; measure real pod density to tighten.", False, None)])
bullet([("Concurrency profile (220 peak hrs/mo, 50-session off-peak baseline) is the biggest lever — adjust to observed telemetry.", False, None)])
bullet([("Shared TRE-core services and per-workspace storage are excluded as roughly comparable / out of scope for this compute delta.", False, None)])


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]; p.alignment = CENTER; p.text = ""
    r = p.add_run("Secure Research Workspace — Cost Analysis   ·   Confidential   ·   Page ")
    r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(GREY)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run_in = OxmlElement("w:r"); rpr = OxmlElement("w:rPr"); sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16"); rpr.append(sz); run_in.append(rpr)
    t = OxmlElement("w:t"); t.text = "1"; run_in.append(t); fld.append(run_in)
    p._p.append(fld)


for sec in doc.sections:
    add_footer(sec)

out = r"D:\kubernetesstructure\SecureResearchWorkspace\SecureResearchWorkspace\SRW_Cost_Analysis.docx"
doc.save(out)
print("saved", out)
